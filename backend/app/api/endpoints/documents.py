"""
Endpoints de Documentos Institucionales (ONB-002).
Soporta dos fuentes: archivos (PDF/DOCX/TXT) y enlaces (URLs).
Extraccion automatica de texto para contexto IA (RAG snapshot).
NINGUN archivo se almacena en disco: solo se persiste el texto extraido en DB.
"""
import io
import os
import re
import tempfile
from datetime import datetime
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from pydantic import BaseModel
from sqlalchemy.orm import Session
from typing import Optional

from app.db.database import get_db
from app.db.models import Document, DOC_SUBTYPES, User, RoleEnum

VALID_DOC_TYPES = set(DOC_SUBTYPES.keys())
from app.core.security import get_current_tenant_id, get_current_user, require_role

router = APIRouter()

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
MAX_EXTRACT_CHARS = 200_000


class DocumentResponse(BaseModel):
    id: int
    source_type: str  # "file" | "link"
    filename: str
    file_type: str
    url: Optional[str] = None
    doc_type: str
    doc_subtype: Optional[str] = None
    char_count: int
    extracted_text: Optional[str] = None
    tenant_id: int
    created_at: datetime

    class Config:
        from_attributes = True


class LinkRequest(BaseModel):
    url: str
    title: str
    doc_type: str = "general"
    doc_subtype: Optional[str] = None


def _doc_to_response(d: Document) -> DocumentResponse:
    return DocumentResponse(
        id=d.id,
        source_type=d.source_type or "file",
        filename=d.filename,
        file_type=d.file_type,
        url=d.url,
        doc_type=d.doc_type or "general",
        doc_subtype=d.doc_subtype,
        char_count=d.char_count or 0,
        extracted_text=d.extracted_text,
        tenant_id=d.tenant_id,
        created_at=d.created_at,
    )


# ── Extraccion de texto (en memoria, sin guardar archivos) ──

def _extract_text_from_bytes(data: bytes, file_type: str) -> str:
    """Extrae texto de bytes en memoria. No escribe nada a disco excepto temps necesarios."""
    try:
        if file_type == "txt":
            return data.decode("utf-8", errors="ignore")[:MAX_EXTRACT_CHARS]

        if file_type == "pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(data))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                    if len(text) > MAX_EXTRACT_CHARS:
                        break
                return text[:MAX_EXTRACT_CHARS]
            except ImportError:
                return "[PyPDF2 no instalado]"

        if file_type == "docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(data))
                text = "\n".join([p.text for p in doc.paragraphs])
                return text[:MAX_EXTRACT_CHARS]
            except ImportError:
                return "[python-docx no instalado]"

    except Exception:
        return ""
    return ""


def _extract_text_from_url(url: str) -> tuple[str, str]:
    """
    Extrae texto de una URL. Retorna (texto, file_type).
    Todo se procesa en memoria, sin guardar archivos.
    """
    import urllib.request
    import urllib.error

    try:
        req = urllib.request.Request(url, headers={
            "User-Agent": "Enfoque365-DocExtractor/1.0",
            "Accept": "text/html,application/pdf,*/*",
        })
        with urllib.request.urlopen(req, timeout=30) as response:
            content_type = response.headers.get("Content-Type", "").lower()
            data = response.read(MAX_FILE_SIZE)

            # PDF remoto -> extraer en memoria
            if "pdf" in content_type or url.lower().endswith(".pdf"):
                text = _extract_text_from_bytes(data, "pdf")
                return text, "pdf"

            # DOCX remoto -> extraer en memoria
            if "wordprocessingml" in content_type or url.lower().endswith(".docx"):
                text = _extract_text_from_bytes(data, "docx")
                return text, "docx"

            # HTML (default) -> limpiar tags
            html = data.decode("utf-8", errors="ignore")
            text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
            text = text.replace("&nbsp;", " ").replace("&quot;", '"')
            return text[:MAX_EXTRACT_CHARS], "html"

    except urllib.error.HTTPError as e:
        return f"[Error HTTP {e.code} al acceder a la URL]", "link"
    except urllib.error.URLError as e:
        return f"[Error de conexion: {e.reason}]", "link"
    except Exception as e:
        return f"[Error al extraer contenido: {str(e)[:200]}]", "link"


def _validate_doc_classification(doc_type: str, doc_subtype: str = "") -> tuple[str, Optional[str]]:
    """Valida tipo y subtipo. Retorna (tipo_validado, subtipo_validado)."""
    validated_type = doc_type if doc_type in VALID_DOC_TYPES else "general"
    valid_subtype = None
    if doc_subtype:
        valid_subtypes = DOC_SUBTYPES.get(validated_type, [])
        if doc_subtype in valid_subtypes:
            valid_subtype = doc_subtype
    return validated_type, valid_subtype


# ── Endpoints ────────────────────────────────

@router.get("/subtypes", response_model=dict[str, list[str]])
def get_subtypes():
    """Retorna el catalogo de subtipos por tipo de documento."""
    return DOC_SUBTYPES


@router.get("/", response_model=list[DocumentResponse])
def list_documents(
    tenant_id: int = Depends(get_current_tenant_id),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Lista todos los documentos del tenant (archivos + enlaces)."""
    docs = db.query(Document).filter(Document.tenant_id == tenant_id).all()
    return [_doc_to_response(d) for d in docs]


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    doc_type: str = Form("general"),
    doc_subtype: str = Form(""),
    tenant_id: int = Depends(get_current_tenant_id),
    user: User = Depends(require_role(RoleEnum.ADMIN, RoleEnum.ESTRATEGA)),
    db: Session = Depends(get_db),
):
    """
    Sube un archivo (PDF, DOCX, TXT). Extrae el texto en memoria
    y lo persiste en la DB. El archivo original NO se almacena en disco.
    """
    ext = file.filename.rsplit(".", 1)[-1].lower() if file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Tipo de archivo no permitido: .{ext}. Permitidos: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Archivo excede el limite de 50 MB")

    # Extraer texto en memoria -- sin guardar archivo a disco
    extracted = _extract_text_from_bytes(content, ext)
    validated_type, valid_subtype = _validate_doc_classification(doc_type, doc_subtype)

    doc = Document(
        source_type="file",
        filename=file.filename,
        file_type=ext,
        file_path=None,
        doc_type=validated_type,
        doc_subtype=valid_subtype,
        extracted_text=extracted,
        char_count=len(extracted),
        tenant_id=tenant_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return _doc_to_response(doc)


@router.post("/link", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
def add_link(
    payload: LinkRequest,
    tenant_id: int = Depends(get_current_tenant_id),
    user: User = Depends(require_role(RoleEnum.ADMIN, RoleEnum.ESTRATEGA)),
    db: Session = Depends(get_db),
):
    """
    Registra un enlace documental. Extrae el contenido de la URL
    automaticamente (snapshot RAG) para alimentar el contexto IA.
    """
    url = payload.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="La URL debe comenzar con http:// o https://")

    extracted, detected_type = _extract_text_from_url(url)
    validated_type, valid_subtype = _validate_doc_classification(
        payload.doc_type, payload.doc_subtype or ""
    )

    doc = Document(
        source_type="link",
        filename=payload.title,
        file_type=detected_type,
        file_path=None,
        url=url,
        doc_type=validated_type,
        doc_subtype=valid_subtype,
        extracted_text=extracted,
        char_count=len(extracted),
        tenant_id=tenant_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return _doc_to_response(doc)


@router.post("/{doc_id}/refresh", response_model=DocumentResponse)
def refresh_link(
    doc_id: int,
    tenant_id: int = Depends(get_current_tenant_id),
    user: User = Depends(require_role(RoleEnum.ADMIN, RoleEnum.ESTRATEGA)),
    db: Session = Depends(get_db),
):
    """Re-extrae el contenido de un enlace documental (actualizar snapshot)."""
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.tenant_id == tenant_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    if doc.source_type != "link" or not doc.url:
        raise HTTPException(status_code=400, detail="Solo se pueden actualizar documentos de tipo enlace")

    extracted, detected_type = _extract_text_from_url(doc.url)
    doc.extracted_text = extracted
    doc.char_count = len(extracted)
    doc.file_type = detected_type
    db.commit()
    db.refresh(doc)
    return _doc_to_response(doc)


@router.delete("/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(
    doc_id: int,
    tenant_id: int = Depends(get_current_tenant_id),
    user: User = Depends(require_role(RoleEnum.ADMIN)),
    db: Session = Depends(get_db),
):
    """Elimina un documento (solo registro en DB, no hay archivo en disco)."""
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.tenant_id == tenant_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    db.delete(doc)
    db.commit()
